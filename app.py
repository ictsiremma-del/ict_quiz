from flask import (Flask, render_template, request, redirect,
    url_for, session, send_file, jsonify)
import json, os, datetime, re, random, threading, shutil, zipfile, socket
from urllib.parse import urlencode
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from werkzeug.utils import secure_filename
from werkzeug.security import generate_password_hash, check_password_hash
import PyPDF2
from database import (
    init_db,
    db_load_results, db_save_result, db_delete_result,
    db_load_questions, db_save_questions,
    db_load_hw, db_save_hw, db_del_hw,
    db_load_bece, db_save_bece,
    db_load_assignments, db_save_assignment, db_delete_assignment,
    db_load_attendance, db_save_attendance,
    db_load_bank, db_add_to_bank, db_delete_from_bank,
    db_register_school, db_get_all_schools, db_get_school,
    db_update_school_status, db_school_code_exists
)

try:
    import requests as req_lib
    REQUESTS_OK = True
except ImportError:
    REQUESTS_OK = False

app = Flask(__name__)
app.secret_key = os.environ.get("SECRET_KEY", "gloriouspearlsquiz2024")
ADMIN_PASSWORD = os.environ.get("ADMIN_PASSWORD", "admin2024")

SCHOOL_NAME        = "Glorious Pearls Complex School"
GROQ_API_KEY       = os.environ.get("GROQ_API_KEY", "")
GROQ_MODEL         = "llama-3.3-70b-versatile"
GROQ_URL           = "https://api.groq.com/openai/v1/chat/completions"
RESULTS_FILE       = "results.json"
HOMEWORK_FILE      = "homework_saves.json"
UPLOAD_FOLDER      = "static/uploads"
ALLOWED_IMG        = {"png","jpg","jpeg","gif","webp"}
ALLOWED_PDF        = {"pdf"}
TEACHERS_FILE      = "teachers.json"
ADMIN_PASSWORD     = os.environ.get("ADMIN_PASSWORD", "admin2024")
GOOGLE_CLIENT_ID   = os.environ.get("GOOGLE_CLIENT_ID", "523848643345-5pjbo9avqt8b8s2gm1rc9u4mrk7cfkcf.apps.googleusercontent.com")
GOOGLE_CLIENT_SECRET     = os.environ.get("GOOGLE_CLIENT_SECRET", "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")
GOOGLE_AUTH_REDIRECT_URI = os.environ.get("GOOGLE_REDIRECT_URI", "https://ictquiz-production.up.railway.app/google-callback")
GOOGLE_AUTH_ENDPOINT     = "https://accounts.google.com/o/oauth2/v2/auth"
GOOGLE_TOKEN_ENDPOINT    = "https://oauth2.googleapis.com/token"
GOOGLE_USERINFO_ENDPOINT = "https://openidconnect.googleapis.com/v1/userinfo"

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs("static", exist_ok=True)
os.makedirs("backups", exist_ok=True)
os.makedirs("templates", exist_ok=True)

def load_custom_teachers():
    if not os.path.exists(TEACHERS_FILE):
        with open(TEACHERS_FILE,"w",encoding="utf-8") as f: json.dump({},f)
        return {}
    try:
        with open(TEACHERS_FILE,"r",encoding="utf-8") as f: return json.load(f) or {}
    except Exception as e:
        print("Unable to load teachers file:", e); return {}

def save_custom_teachers(data):
    try:
        with open(TEACHERS_FILE,"w",encoding="utf-8") as f: json.dump(data,f,indent=2)
    except Exception as e:
        print("Unable to save teachers file:", e)

CUSTOM_TEACHERS = load_custom_teachers()

try: init_db()
except Exception as e: print("DB init warning:", e)

app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
app.config["MAX_CONTENT_LENGTH"] = 16 * 1024 * 1024
_lock = threading.Lock()

TEACHERS = {
    "sir_emma":      {"name":"Sir Emma",      "password":"Letgohome",  "subjects":["ICT"],                                                                                                                                                "is_head":True,  "managed_class":"JHS 2"},
    "sir_eddie":     {"name":"Sir Eddie",     "password":"eddie123",   "subjects":["Science"],                                                                                                                                            "is_head":False, "managed_class":"Basic 6"},
    "sir_bismark":   {"name":"Sir Bismark",   "password":"bismark123", "subjects":["Creative Arts and Design","Career Technology"],                                                                                                        "is_head":False, "managed_class":"Basic 5"},
    "sir_sam":       {"name":"Sir Sam",       "password":"sam123",     "subjects":["French"],                                                                                                                                             "is_head":False, "managed_class":"JHS 1"},
    "ms_gasu":       {"name":"Ms. Gasu",      "password":"gasu123",    "subjects":["English Language","Religious and Moral Education"],                                                                                                    "is_head":False, "managed_class":None},
    "sir_sackey":    {"name":"Sir Sackey",    "password":"sackey123",  "subjects":["ICT","Mathematics","Science","English Language","Social Studies","Religious and Moral Education","French","Career Technology","Creative Arts and Design"],"is_head":False, "managed_class":"Basic 2"},
    "master":        {"name":"Master",        "password":"master123",  "subjects":["Mathematics"],                                                                                                                                        "is_head":False, "managed_class":"JHS 3"},
    "sir_otoo":      {"name":"Sir Otoo",      "password":"otoo123",    "subjects":["Ghanaian Language"],                                                                                                                                  "is_head":False, "managed_class":None},
    "mr_tinkorange": {"name":"Mr. Tinkorange","password":"Sir TK",     "subjects":["Mathematics"],                                                                                                                                        "is_head":False, "managed_class":"Basic 4A"},
    "mrs_victoria":  {"name":"Mrs. Victoria", "password":"madvic",     "subjects":["English Language","Religious and Moral Education"],                                                                                                    "is_head":False, "managed_class":"Basic 4B"},
    "ms_priscilla":  {"name":"Ms. Priscilla", "password":"priscilla",  "subjects":["ICT","Mathematics","Science","English Language","Social Studies","Religious and Moral Education","French","Career Technology","Creative Arts and Design"],"is_head":False, "managed_class":"Basic 1"},
    "mr_edem":       {"name":"Mr. Edem",      "password":"siredem",    "subjects":["ICT","Mathematics","Science","English Language","Social Studies","Religious and Moral Education","French","Career Technology","Creative Arts and Design"],"is_head":False, "managed_class":"Basic 3"},
}

SUBJECTS = ["ICT","Mathematics","Science","English Language","Social Studies",
            "Religious and Moral Education","French","Ghanaian Language",
            "Career Technology","Creative Arts and Design"]

CLASSES = ["Basic 1","Basic 2","Basic 3","Basic 4A","Basic 4B",
           "Basic 5","Basic 6","JHS 1","JHS 2","JHS 3"]

GRADES = [(90,100,"A","Excellent"),(80,89,"B","Very Good"),(70,79,"C","Good"),
          (60,69,"D","Average"),(50,59,"E","Below Average"),(0,49,"F","Fail")]

SUBJECT_COLORS = {
    "ICT":"#2e86ab","Mathematics":"#e74c3c","Science":"#27ae60",
    "English Language":"#8e44ad","Social Studies":"#f39c12",
    "Religious and Moral Education":"#16a085","French":"#2980b9",
    "Ghanaian Language":"#d35400","Career Technology":"#c0392b",
    "Creative Arts and Design":"#e91e8c"}

SUBJECT_ICONS = {
    "ICT":"💻","Mathematics":"🔢","Science":"🔬","English Language":"📚",
    "Social Studies":"🌍","Religious and Moral Education":"✝️","French":"🇫🇷",
    "Ghanaian Language":"🇬🇭","Career Technology":"⚙️","Creative Arts and Design":"🎨"}

ASSESSMENT_TYPES = {
    "class_test":{"label":"Class Test","icon":"📝","timer_minutes":30,"show_answers_after":True,"can_pause":False,"shuffle":True,"description":"Timed 30-minute test"},
    "exam":      {"label":"Examination","icon":"📋","timer_minutes":60,"show_answers_after":False,"can_pause":False,"shuffle":True,"description":"Timed 60-minute exam"},
    "homework":  {"label":"Homework","icon":"🏠","timer_minutes":0,"show_answers_after":True,"can_pause":True,"shuffle":False,"description":"No timer — save and continue later"},
}

CURRICULUM_STRANDS = {
    "ICT":{"Computing Systems":["Hardware Components","Software and Operating Systems","Input and Output Devices","Storage Devices","Computer Networks"],"Data and Information":["Data Collection and Processing","Spreadsheets","Databases","Information Representation","Data Security"],"Algorithms and Programming":["Problem Solving","Algorithms and Flowcharts","Scratch Programming","Python Basics","Web Design"],"Communication and Collaboration":["Internet and Email","Social Media Safety","Digital Citizenship","Online Collaboration Tools"],"Digital Literacy":["Word Processing","Presentation Software","Digital Safety and Ethics","Emerging Technologies"]},
    "Mathematics":{"Number and Numeration":["Whole Numbers","Fractions and Decimals","Percentages and Ratios","Indices and Logarithms","Number Bases"],"Algebra":["Algebraic Expressions","Linear Equations","Quadratic Equations","Inequalities","Sequences and Series"],"Geometry and Measurement":["Shapes and Properties","Angles and Lines","Perimeter Area and Volume","Pythagoras Theorem","Trigonometry"],"Data Handling":["Data Collection","Statistical Graphs","Measures of Central Tendency","Probability"],"Patterns and Relations":["Number Patterns","Functions and Relations","Matrices"]},
    "Science":{"Diversity of Matter":["Living and Non-living Things","Classification of Living Things","Properties of Matter","States of Matter","Mixtures and Solutions"],"Cycles":["Life Cycles","Water Cycle","Carbon Cycle","Rock Cycle","Nutrient Cycle"],"Systems and Interactions":["Human Body Systems","Ecosystems","Food Chains and Webs","Digestive System","Reproductive System"],"Forces and Energy":["Types of Forces","Energy and its Forms","Electricity","Light and Sound","Heat Energy"],"Earth and Space":["Solar System","Earths Structure","Weather and Climate","Natural Disasters","Environmental Issues"]},
    "English Language":{"Reading":["Comprehension","Vocabulary in Context","Inference and Deduction","Literary Devices"],"Writing":["Narrative Writing","Expository Writing","Argumentative Writing","Letter Writing","Summary Writing"],"Grammar and Usage":["Parts of Speech","Tenses","Sentence Structure","Punctuation","Concord and Agreement"],"Listening and Speaking":["Oral Communication","Phonics and Pronunciation","Debate and Discussion","Story Telling"],"Literature":["Poetry","Prose","Drama","African Literature"]},
    "Social Studies":{"Citizenship":["Rights and Responsibilities","Democracy and Governance","National Symbols","Rule of Law","Civic Participation"],"Geography and Environment":["Map Reading and Skills","Regions of Ghana","Natural Resources","Environmental Conservation","Climate and Vegetation"],"History and Culture":["Pre-Colonial Ghana","Colonial Period","Independence of Ghana","Ghanaian Culture and Traditions","African History"],"Economics and Development":["Economic Activities","Trade and Commerce","Agriculture in Ghana","Infrastructure and Development","Global Issues"]},
    "Religious and Moral Education":{"Religious Beliefs":["Christianity","Islam","African Traditional Religion","Religious Tolerance","Festivals and Celebrations"],"Moral Values":["Honesty and Integrity","Respect and Responsibility","Love and Compassion","Justice and Fairness","Forgiveness"],"Family and Society":["Family Life","Community Living","Social Relationships","Marriage and Parenthood"],"Health and Personal Development":["Personal Hygiene","Avoiding Bad Habits","Peer Pressure","Life Skills"]},
    "French":{"Listening and Speaking":["Greetings and Introductions","Numbers and Colours","Days Months and Seasons","Oral Conversations","French Pronunciation"],"Reading":["Reading Comprehension","French Texts","Vocabulary Building"],"Writing":["Sentence Construction","Letter Writing in French","Descriptive Writing"],"Grammar":["Articles and Nouns","Verbs and Conjugation","Adjectives and Adverbs","Pronouns","Prepositions"],"Culture":["French Speaking Countries","French Culture and Traditions","Francophone Africa"]},
    "Ghanaian Language":{"Oral Language":["Greetings and Conversations","Story Telling","Proverbs and Idioms","Oral Traditions"],"Reading":["Reading Fluency","Comprehension","Vocabulary"],"Writing":["Letter Writing","Sentence Construction","Creative Writing"],"Grammar":["Nouns and Pronouns","Verbs and Tenses","Sentence Structure","Punctuation"],"Culture and Values":["Ghanaian Traditions","Festivals","Chieftaincy","Folklore"]},
    "Career Technology":{"Home Economics":["Food and Nutrition","Clothing and Textiles","Home Management","Child Development"],"Technical Drawing":["Drawing Instruments","Geometrical Constructions","Orthographic Projection","Isometric Drawing"],"Workshop Technology":["Safety in the Workshop","Woodwork","Metalwork","Electrical Work"],"Agricultural Science":["Crop Production","Animal Husbandry","Farm Tools and Equipment","Soil and Fertilizers"],"Business Studies":["Entrepreneurship","Book Keeping","Office Practice","Buying and Selling"]},
    "Creative Arts and Design":{"Visual Arts":["Drawing and Painting","Sculpture and Modelling","Graphic Design","Printmaking","Weaving and Basketry"],"Performing Arts":["Music and Rhythm","Dance","Drama and Theatre","Traditional Performances"],"Design and Technology":["Product Design","Fashion and Textiles","Pottery and Ceramics","Jewellery Making"],"Ghanaian Cultural Arts":["Kente Weaving","Adinkra Symbols","Traditional Music","Ghanaian Dance Forms"]},
}

BECE_SUBJECTS = ["ICT","Mathematics","Science","English Language","Social Studies",
                 "Religious and Moral Education","French","Career Technology","Creative Arts and Design"]
BECE_YEARS = [str(y) for y in range(2010,2026)]
GHANA_REGIONS = ["Ahafo","Ashanti","Bono","Bono East","Central","Eastern","Greater Accra",
                 "North East","Northern","Oti","Savannah","Upper East","Upper West","Volta","Western","Western North"]

def get_local_ip():
    try:
        s=socket.socket(socket.AF_INET,socket.SOCK_DGRAM); s.connect(("8.8.8.8",80))
        ip=s.getsockname()[0]; s.close(); return ip
    except: return "127.0.0.1"

def load_qs(s): return db_load_questions(s)
def save_qs(s,qs): db_save_questions(s,qs)
def load_results(): return db_load_results()
def save_result(r): db_save_result(r)
def load_hw(): return db_load_hw()
def save_hw(k,d): db_save_hw(k,d)
def del_hw(k): db_del_hw(k)

def get_grade(p):
    for lo,hi,g,r in GRADES:
        if lo<=p<=hi: return g,r
    return "F","Fail"

def gcolor(g): return {"A":"#27ae60","B":"#2980b9","C":"#8e44ad","D":"#f39c12","E":"#e67e22","F":"#e74c3c"}.get(g,"#95a5a6")
def allowed_img(fn): return "." in fn and fn.rsplit(".",1)[1].lower() in ALLOWED_IMG
def allowed_pdf(fn): return "." in fn and fn.rsplit(".",1)[1].lower() in ALLOWED_PDF

def keyword_score(ans,model,marks):
    if not ans or not model: return 0,[]
    kws=list(set([w.strip().lower() for w in model.split() if len(w.strip())>3]))
    matched=[k for k in kws if k in ans.lower()]
    if not kws: return marks,[]
    return min(round(len(matched)/len(kws)*marks),marks),matched

def get_qs_for_student(subject,cls,shuffle=True,strand=None,sub_strand=None):
    qs=load_qs(subject); filtered=[]
    for q in qs:
        ac=q.get("assigned_classes",[])
        if ac and cls not in ac: continue
        if strand and q.get("strand","")!=strand: continue
        if sub_strand and q.get("sub_strand","")!=sub_strand: continue
        filtered.append(dict(q))
    if shuffle:
        random.shuffle(filtered)
        for q in filtered:
            if q["type"]=="mcq" and q.get("options"):
                opts=q["options"].copy(); random.shuffle(opts); q["options"]=opts
    return filtered

def current_teacher(): return session.get("teacher_data")

def get_teacher_by_id(uid):
    if not uid: return None, None
    if uid in TEACHERS: return uid, TEACHERS[uid]
    if uid in CUSTOM_TEACHERS: return uid, CUSTOM_TEACHERS[uid]
    lower = uid.lower()
    for key, teacher in {**TEACHERS, **CUSTOM_TEACHERS}.items():
        if teacher.get("email","").lower() == lower: return key, teacher
    return None, None

def teacher_exists(uid, email=None):
    _, teacher = get_teacher_by_id(uid)
    if teacher: return True
    if email:
        lower = email.lower()
        for teacher in {**TEACHERS, **CUSTOM_TEACHERS}.values():
            if teacher.get("email","").lower() == lower: return True
    return False

def can_access(subject):
    t=current_teacher()
    if not t: return False
    return t.get("is_head") or subject in t.get("subjects",[])

def verify_password(stored, provided):
    if not stored: return False
    if stored.startswith("pbkdf2:") or stored.startswith("scrypt:"):
        return check_password_hash(stored, provided)
    return stored == provided

def extract_pdf(filepath):
    qs=[]
    try:
        reader=PyPDF2.PdfReader(open(filepath,"rb"))
        text="\n".join(p.extract_text() or "" for p in reader.pages)
        lines=[l.strip() for l in text.split("\n") if l.strip()]
        i=0; qid=1
        while i<len(lines):
            ln=lines[i]
            m=re.match(r"^(\d+)[.)]\s+(.+)$",ln)
            if m:
                qt=m.group(2); opts=[]; j=i+1
                while j<len(lines) and j<=i+6:
                    om=re.match(r"^[A-Da-d][.)]\s*(.+)$",lines[j])
                    if om: opts.append(om.group(1).strip()); j+=1
                    else: break
                if len(opts)>=2:
                    ans=""
                    if j<len(lines):
                        am=re.match(r"^(?:Answer|ANS)[:\s]+(.+)$",lines[j],re.I)
                        if am: ans=am.group(1).strip(); j+=1
                    qs.append({"id":qid,"type":"mcq","question":qt,"options":opts[:4],"answer":ans,"marks":2,"image":"","assigned_classes":[],"strand":"","sub_strand":""}); qid+=1; i=j; continue
            m2=re.match(r"^(?:TF|T/F)[:\s]+(?:\d+[.)]\s*)?(.+)$",ln,re.I)
            if m2:
                q={"id":qid,"type":"tf","question":m2.group(1),"answer":"True","marks":1,"image":"","assigned_classes":[],"strand":"","sub_strand":""}
                if i+1<len(lines):
                    am=re.match(r"^(?:Answer|ANS)[:\s]+(True|False)$",lines[i+1],re.I)
                    if am: q["answer"]=am.group(1).capitalize(); i+=1
                qs.append(q); qid+=1; i+=1; continue
            m3=re.match(r"^(?:FB|BLANK)[:\s]+(?:\d+[.)]\s*)?(.+)$",ln,re.I)
            if m3:
                ans=""
                if i+1<len(lines):
                    am=re.match(r"^(?:Answer|ANS)[:\s]+(.+)$",lines[i+1],re.I)
                    if am: ans=am.group(1).strip(); i+=1
                qs.append({"id":qid,"type":"fitb","question":m3.group(1),"answer":ans,"marks":2,"image":"","assigned_classes":[],"strand":"","sub_strand":""}); qid+=1; i+=1; continue
            m4=re.match(r"^(?:TH|THEORY)[:\s]+(?:\d+[.)]\s*)?(.+)$",ln,re.I)
            if m4:
                qt=m4.group(1); model=""; marks=5
                mk=re.search(r"\[(\d+)\s*marks?\]",qt,re.I)
                if mk: marks=int(mk.group(1))
                if i+1<len(lines):
                    am=re.match(r"^(?:Model|Answer|ANS)[:\s]+(.+)$",lines[i+1],re.I)
                    if am: model=am.group(1).strip(); i+=1
                qs.append({"id":qid,"type":"theory","question":qt,"model_answer":model,"marks":marks,"image":"","assigned_classes":[],"strand":"","sub_strand":""}); qid+=1; i+=1; continue
            i+=1
    except Exception as e: print("PDF error:",e)
    return qs

def load_bece(s,y): return db_load_bece(s,y)
def save_bece(s,y,qs): db_save_bece(s,y,qs)

def bece_years_for(subject):
    safe=subject.replace(" ","_"); years=[]
    for f in os.listdir("."):
        if f.startswith("bece_{}_".format(safe)) and f.endswith(".json"):
            yr=f.replace("bece_{}_".format(safe),"").replace(".json","")
            qs=load_bece(subject,yr)
            if qs: years.append({"year":yr,"count":len(qs)})
    return sorted(years,key=lambda x:x["year"],reverse=True)

@app.route("/")
def landing():
    return render_template("landing.html",school=SCHOOL_NAME)

@app.route("/student")
def index():
    hw_key=session.get("hw_save_key"); hw_save=load_hw().get(hw_key) if hw_key else None
    return render_template("index.html",classes=CLASSES,subjects=SUBJECTS,school=SCHOOL_NAME,
        subject_colors=SUBJECT_COLORS,subject_icons=SUBJECT_ICONS,assessment_types=ASSESSMENT_TYPES,
        curriculum_strands=CURRICULUM_STRANDS,hw_save=hw_save,hw_key=hw_key)

@app.route("/start",methods=["POST"])
def start():
    name=request.form.get("name","").strip(); cls=request.form.get("class","").strip()
    subject=request.form.get("subject","").strip(); atype=request.form.get("assessment_type","class_test")
    strand=request.form.get("strand","").strip(); sub_strand=request.form.get("sub_strand","").strip()
    if not name or not cls or not subject: return redirect(url_for("index"))
    aconf=ASSESSMENT_TYPES.get(atype,ASSESSMENT_TYPES["class_test"])
    qs=get_qs_for_student(subject,cls,aconf["shuffle"],strand or None,sub_strand or None)
    if not qs: return render_template("no_questions.html",school=SCHOOL_NAME,subject=subject,cls=cls)
    session.update({"student_name":name,"student_class":cls,"student_subject":subject,
        "assessment_type":atype,"aconf":aconf,"start_time":datetime.datetime.now().isoformat(),
        "questions":qs,"strand":strand,"sub_strand":sub_strand})
    return render_template("quiz.html",questions=qs,student_name=name,student_class=cls,subject=subject,
        atype=atype,aconf=aconf,subject_color=SUBJECT_COLORS.get(subject,"#2e86ab"),
        subject_icon=SUBJECT_ICONS.get(subject,"📝"),timer=aconf["timer_minutes"]*60,
        school=SCHOOL_NAME,can_pause=aconf["can_pause"],saved_answers={})

@app.route("/save_homework",methods=["POST"])
def save_homework():
    name=session.get("student_name",""); subject=session.get("student_subject","")
    if not name or not subject: return jsonify({"ok":False})
    cls=session.get("student_class","")
    key="{}_{}_{}_hw".format(name.lower().replace(" ","_"),cls.lower().replace(" ","_"),subject.lower().replace(" ","_"))
    save_hw(key,{"name":name,"class":cls,"subject":subject,"answers":request.json.get("answers",{}),
        "questions":session.get("questions",[]),"aconf":session.get("aconf",{}),
        "saved_at":datetime.datetime.now().strftime("%d/%m/%Y %H:%M")})
    session["hw_save_key"]=key
    return jsonify({"ok":True})

@app.route("/resume_homework/<key>")
def resume_homework(key):
    save=load_hw().get(key)
    if not save: return redirect(url_for("index"))
    aconf=save.get("aconf",ASSESSMENT_TYPES["homework"])
    session.update({"student_name":save["name"],"student_class":save["class"],
        "student_subject":save["subject"],"assessment_type":"homework","aconf":aconf,
        "questions":save["questions"],"start_time":datetime.datetime.now().isoformat()})
    return render_template("quiz.html",questions=save["questions"],student_name=save["name"],
        student_class=save["class"],subject=save["subject"],atype="homework",aconf=aconf,
        subject_color=SUBJECT_COLORS.get(save["subject"],"#2e86ab"),
        subject_icon=SUBJECT_ICONS.get(save["subject"],"📝"),
        timer=0,school=SCHOOL_NAME,can_pause=True,saved_answers=save.get("answers",{}))

@app.route("/submit",methods=["POST"])
def submit():
    qs=session.get("questions",[]); name=session.get("student_name","Unknown")
    cls=session.get("student_class","Unknown"); subject=session.get("student_subject","Unknown")
    atype=session.get("assessment_type","class_test"); aconf=session.get("aconf",ASSESSMENT_TYPES["class_test"])
    strand=session.get("strand",""); sub_strand=session.get("sub_strand","")
    total=sum(q["marks"] for q in qs); earned=0; details=[]
    for q in qs:
        qid=str(q["id"]); qtype=q["type"]; given=request.form.get("q{}".format(qid),"").strip()
        correct=q.get("answer","").strip(); marks=q["marks"]
        if qtype in ["mcq","tf","fitb"]:
            if qtype=="mcq":
                opts=q.get("options",[]); letters=["A","B","C","D"]
                given_letter=given.upper() if given.upper() in letters else ""
                if not given_letter:
                    for i,opt in enumerate(opts):
                        if opt.strip().lower()==given.strip().lower() and i<4: given_letter=letters[i]; break
                correct_letter=correct.upper() if correct.upper() in letters else ""
                if not correct_letter:
                    for i,opt in enumerate(opts):
                        if opt.strip().lower()==correct.strip().lower() and i<4: correct_letter=letters[i]; break
                right=given_letter==correct_letter if given_letter and correct_letter else False
                given_display=("{}: {}".format(given_letter,opts[letters.index(given_letter)]) if given_letter in letters and letters.index(given_letter)<len(opts) else given_letter) if given_letter else "Not answered"
                correct_display=("{}: {}".format(correct_letter,opts[letters.index(correct_letter)]) if correct_letter in letters and letters.index(correct_letter)<len(opts) else correct_letter) if correct_letter else correct
            else:
                right=given.lower()==correct.lower(); given_display=given or "Not answered"; correct_display=correct
            got=marks if right else 0; earned+=got
            details.append({"question":q["question"],"type":qtype,"given":given_display,"correct":correct_display,"correct_flag":right,"marks_earned":got,"marks_total":marks,"strand":q.get("strand",""),"image":q.get("image","")})
        elif qtype=="theory":
            got,matched=keyword_score(given,q.get("model_answer",""),marks); earned+=got
            details.append({"question":q["question"],"type":"theory","given":given or "Not answered","correct":q.get("model_answer",""),"correct_flag":got>=marks*0.5,"marks_earned":got,"marks_total":marks,"matched_keywords":matched,"strand":q.get("strand",""),"image":q.get("image","")})
        elif qtype=="label":
            labels=q.get("labels",[]); qe=0; ld=[]
            for lbl in labels:
                ans=request.form.get("q{}_label_{}".format(qid,lbl["id"]),"").strip()
                right=ans.lower()==lbl["answer"].lower()
                if right: qe+=marks/len(labels) if labels else 0
                ld.append({"label_id":lbl["id"],"given":ans or "Not answered","correct":lbl["answer"],"correct_flag":right})
            qe=round(qe); earned+=qe
            details.append({"question":q["question"],"type":"label","given":"(labelling)","correct":"(see labels)","correct_flag":qe>=marks*0.5,"marks_earned":qe,"marks_total":marks,"label_details":ld,"image":q.get("image","")})
    pct=round((earned/total)*100,1) if total>0 else 0
    grade,remark=get_grade(pct)
    result={"name":name,"class":cls,"subject":subject,"assessment_type":atype,"assessment_label":aconf.get("label","Class Test"),"strand":strand,"sub_strand":sub_strand,"score":earned,"total":total,"percentage":pct,"grade":grade,"remark":remark,"date":datetime.datetime.now().strftime("%d/%m/%Y"),"time":datetime.datetime.now().strftime("%H:%M"),"details":details}
    save_result(result)
    if atype=="homework":
        hw_key=session.get("hw_save_key")
        if hw_key: del_hw(hw_key); session.pop("hw_save_key",None)
    return render_template("result.html",result=result,school=SCHOOL_NAME,grade_color=gcolor(grade),
        subject_color=SUBJECT_COLORS.get(subject,"#2e86ab"),subject_icon=SUBJECT_ICONS.get(subject,"📝"),
        show_answers=aconf.get("show_answers_after",True))

@app.route("/teacher",methods=["GET","POST"])
def teacher():
    if request.method=="POST":
        uid=request.form.get("username","").strip(); pwd=request.form.get("password","").strip()
        # Reload custom teachers fresh on every login attempt
        fresh=load_custom_teachers(); CUSTOM_TEACHERS.update(fresh)
        stored_uid, t = get_teacher_by_id(uid)
        if t:
            status=t.get("status","approved")
            if status=="pending": return render_template("teacher_login.html",error="Your registration is pending admin approval.",school=SCHOOL_NAME)
            if status=="rejected": return render_template("teacher_login.html",error="Your registration was rejected. Contact the administrator.",school=SCHOOL_NAME)
            if verify_password(t.get("password",""),pwd):
                session["teacher"]=True; session["teacher_data"]={**t,"uid":stored_uid}
                return redirect(url_for("dashboard"))
        return render_template("teacher_login.html",error="Wrong username or password.",school=SCHOOL_NAME)
    return render_template("teacher_login.html",school=SCHOOL_NAME,error=request.args.get("error"))

@app.route("/teacher_register",methods=["GET","POST"])
def teacher_register():
    if request.method=="POST":
        name=request.form.get("name","").strip(); uid=request.form.get("uid","").strip().lower()
        email=request.form.get("email","").strip().lower(); password=request.form.get("password","")
        confirm=request.form.get("confirm_password",""); role=request.form.get("role","").strip()
        managed_class=request.form.get("managed_class","").strip()
        subjects=request.form.getlist("subjects")
        def err(msg): return render_template("teacher_register.html",school=SCHOOL_NAME,subjects=SUBJECTS,classes=CLASSES,error=msg)
        if not (name and uid and email and password and confirm and role): return err("Please fill in all required fields.")
        if not subjects: return err("Please select at least one subject you teach.")
        if role not in ["subject_teacher","class_manager"]: return err("Please select a valid role.")
        if role=="class_manager" and not managed_class: return err("Class teachers must select their class.")
        if managed_class and managed_class not in CLASSES: return err("Invalid class selected.")
        if password!=confirm: return err("Passwords do not match.")
        if len(password)<6: return err("Password must be at least 6 characters.")
        fresh=load_custom_teachers(); CUSTOM_TEACHERS.update(fresh)
        if teacher_exists(uid,email): return err("A teacher with that username or email already exists.")
        if any(s not in SUBJECTS for s in subjects): return err("One or more selected subjects are invalid.")
        CUSTOM_TEACHERS[uid]={"name":name,"password":generate_password_hash(password),"subjects":subjects,"is_head":False,"role":role,"managed_class":managed_class if managed_class else None,"email":email,"status":"pending"}
        save_custom_teachers(CUSTOM_TEACHERS)
        return redirect(url_for("teacher",error="Registration submitted and is pending admin approval."))
    return render_template("teacher_register.html",school=SCHOOL_NAME,subjects=SUBJECTS,classes=CLASSES,error=None)

@app.route("/google-login")
def google_login():
    return render_template("google_login.html",school=SCHOOL_NAME)

@app.route("/google-auth")
def google_auth():
    if not (GOOGLE_CLIENT_ID and GOOGLE_CLIENT_SECRET and REQUESTS_OK): return redirect(url_for("google_login"))
    params={"client_id":GOOGLE_CLIENT_ID,"redirect_uri":GOOGLE_AUTH_REDIRECT_URI,"response_type":"code","scope":"openid email profile","access_type":"offline","prompt":"select_account"}
    return redirect("{}?{}".format(GOOGLE_AUTH_ENDPOINT,urlencode(params)))

@app.route("/google-callback")
def google_callback():
    error=request.args.get("error")
    if error: return redirect(url_for("teacher",error="Google login cancelled."))
    code=request.args.get("code")
    if not code or not (GOOGLE_CLIENT_ID and GOOGLE_CLIENT_SECRET and REQUESTS_OK): return redirect(url_for("teacher",error="Google login failed."))
    token_resp=req_lib.post(GOOGLE_TOKEN_ENDPOINT,data={"code":code,"client_id":GOOGLE_CLIENT_ID,"client_secret":GOOGLE_CLIENT_SECRET,"redirect_uri":GOOGLE_AUTH_REDIRECT_URI,"grant_type":"authorization_code"},headers={"Content-Type":"application/x-www-form-urlencoded"})
    if token_resp.status_code!=200: return redirect(url_for("teacher",error="Unable to exchange Google code."))
    access_token=token_resp.json().get("access_token")
    if not access_token: return redirect(url_for("teacher",error="Google token not received."))
    profile_resp=req_lib.get(GOOGLE_USERINFO_ENDPOINT,headers={"Authorization":"Bearer {}".format(access_token)})
    if profile_resp.status_code!=200: return redirect(url_for("teacher",error="Unable to fetch Google profile."))
    profile=profile_resp.json(); email=profile.get("email",""); uid=email.split("@")[0] if email else None
    fresh=load_custom_teachers(); CUSTOM_TEACHERS.update(fresh)
    stored_uid,t=get_teacher_by_id(uid or email)
    if t:
        status=t.get("status","approved")
        if status=="pending": return redirect(url_for("teacher",error="Your Google account is pending admin approval."))
        if status=="rejected": return redirect(url_for("teacher",error="Your Google registration was rejected."))
        session["teacher"]=True; session["teacher_data"]={**t,"uid":stored_uid}
        return redirect(url_for("dashboard"))
    session["google_pending"]={"name":profile.get("name","Google User"),"email":email,"uid_suggestion":email.split("@")[0] if email else ""}
    return redirect(url_for("google_complete_registration"))

@app.route("/google-complete-registration",methods=["GET","POST"])
def google_complete_registration():
    pending=session.get("google_pending")
    if not pending: return redirect(url_for("teacher"))
    if request.method=="POST":
        uid=request.form.get("uid","").strip().lower(); role=request.form.get("role","").strip()
        managed_class=request.form.get("managed_class","").strip(); subjects=request.form.getlist("subjects")
        def err(msg): return render_template("google_complete_registration.html",school=SCHOOL_NAME,subjects=SUBJECTS,classes=CLASSES,pending=pending,error=msg)
        if not uid: return err("Please enter a username.")
        if role not in ["subject_teacher","class_manager"]: return err("Please select a valid role.")
        if not subjects: return err("Please select at least one subject you teach.")
        if role=="class_manager" and not managed_class: return err("Class teachers must select their class.")
        if any(s not in SUBJECTS for s in subjects): return err("One or more selected subjects are invalid.")
        fresh=load_custom_teachers(); CUSTOM_TEACHERS.update(fresh)
        if teacher_exists(uid,pending["email"]): return err("That username or email is already registered.")
        CUSTOM_TEACHERS[uid]={"name":pending["name"],"password":"","subjects":subjects,"is_head":False,"role":role,"managed_class":managed_class if managed_class else None,"email":pending["email"],"status":"pending"}
        save_custom_teachers(CUSTOM_TEACHERS)
        session.pop("google_pending",None)
        return redirect(url_for("teacher",error="Registration submitted and is pending admin approval."))
    return render_template("google_complete_registration.html",school=SCHOOL_NAME,subjects=SUBJECTS,classes=CLASSES,pending=pending,error=None)

@app.route("/teacher_logout")
def teacher_logout(): session.clear(); return redirect(url_for("index"))

@app.route("/dashboard")
def dashboard():
    if not session.get("teacher"): return redirect(url_for("teacher"))
    t=current_teacher()
    my_subjects=SUBJECTS if t.get("is_head") else (t.get("subjects") or [])
    sel_sub=request.args.get("subject",my_subjects[0] if my_subjects else "")
    if sel_sub not in my_subjects and my_subjects: sel_sub=my_subjects[0]
    results=load_results()
    if not t.get("is_head"): results=[r for r in results if r.get("subject","") in my_subjects]
    qs=load_qs(sel_sub) if sel_sub else []
    total=len(results); avg=round(sum(r["percentage"] for r in results)/len(results),1) if results else 0
    gc={}
    for r in results: gc[r["grade"]]=gc.get(r["grade"],0)+1
    sc={}
    for r in results: sc[r.get("subject","")]=sc.get(r.get("subject",""),0)+1
    return render_template("dashboard.html",results=results,questions=qs,selected_subject=sel_sub,
        subjects=my_subjects,all_subjects=SUBJECTS,subject_colors=SUBJECT_COLORS,subject_icons=SUBJECT_ICONS,
        total_students=total,avg_score=avg,grade_counts=gc,sub_counts=sc,school=SCHOOL_NAME,
        classes=CLASSES,pdf_msg=session.pop("pdf_msg",None),restore_msg=session.pop("restore_msg",None),
        local_ip=get_local_ip(),teacher=t,is_head=t.get("is_head",False),
        assessment_types=ASSESSMENT_TYPES,strands=CURRICULUM_STRANDS.get(sel_sub,{}),
        curriculum_strands=CURRICULUM_STRANDS)

@app.route("/add_question",methods=["POST"])
def add_question():
    if not session.get("teacher"): return redirect(url_for("teacher"))
    subject=request.form.get("subject","ICT")
    if not can_access(subject): return redirect(url_for("dashboard"))
    qtype=request.form.get("type","mcq"); question=request.form.get("question","").strip()
    marks=int(request.form.get("marks",2)); assigned=request.form.getlist("assigned_classes")
    strand=request.form.get("strand","").strip(); sub_strand=request.form.get("sub_strand","").strip()
    qs=load_qs(subject); new_id=max((q["id"] for q in qs),default=0)+1
    img_path=""
    if "image" in request.files:
        img=request.files["image"]
        if img and img.filename and allowed_img(img.filename):
            fn=secure_filename("{}_{}".format(new_id,img.filename)); img.save(os.path.join(UPLOAD_FOLDER,fn)); img_path="uploads/{}".format(fn)
    base={"id":new_id,"type":qtype,"question":question,"marks":marks,"image":img_path,"assigned_classes":assigned,"strand":strand,"sub_strand":sub_strand}
    if qtype=="mcq": base["options"]=[request.form.get("opt{}".format(i),"").strip() for i in range(1,5)]; base["answer"]=request.form.get("answer","").strip()
    elif qtype=="tf": base["answer"]=request.form.get("tf_answer","True")
    elif qtype=="fitb": base["answer"]=request.form.get("fitb_answer","").strip()
    elif qtype=="theory": base["model_answer"]=request.form.get("model_answer","").strip()
    elif qtype=="label":
        labels=[]
        for li in range(1,7):
            lt=request.form.get("label{}_text".format(li),"").strip(); la=request.form.get("label{}_answer".format(li),"").strip()
            if lt and la: labels.append({"id":li,"text":lt,"answer":la,"x":float(request.form.get("label{}_x".format(li),50)),"y":float(request.form.get("label{}_y".format(li),50))})
        base["labels"]=labels
    qs.append(base); save_qs(subject,qs)
    return redirect(url_for("dashboard",subject=subject))

@app.route("/edit_question/<subject>/<int:qid>",methods=["GET","POST"])
def edit_question(subject,qid):
    if not session.get("teacher"): return redirect(url_for("teacher"))
    if not can_access(subject): return redirect(url_for("dashboard"))
    qs=load_qs(subject); q=next((x for x in qs if x["id"]==qid),None)
    if not q: return redirect(url_for("dashboard",subject=subject))
    if request.method=="POST":
        q["question"]=request.form.get("question","").strip(); q["marks"]=int(request.form.get("marks",q["marks"]))
        q["assigned_classes"]=request.form.getlist("assigned_classes")
        q["strand"]=request.form.get("strand","").strip(); q["sub_strand"]=request.form.get("sub_strand","").strip()
        if q["type"]=="mcq": q["options"]=[request.form.get("opt{}".format(i),"").strip() for i in range(1,5)]; q["answer"]=request.form.get("answer","").strip()
        elif q["type"]=="tf": q["answer"]=request.form.get("tf_answer","True")
        elif q["type"]=="fitb": q["answer"]=request.form.get("fitb_answer","").strip()
        elif q["type"]=="theory": q["model_answer"]=request.form.get("model_answer","").strip()
        if "image" in request.files:
            img=request.files["image"]
            if img and img.filename and allowed_img(img.filename):
                fn=secure_filename("{}_{}".format(qid,img.filename)); img.save(os.path.join(UPLOAD_FOLDER,fn)); q["image"]="uploads/{}".format(fn)
        save_qs(subject,qs); return redirect(url_for("dashboard",subject=subject))
    return render_template("edit_question.html",q=q,subject=subject,school=SCHOOL_NAME,
        classes=CLASSES,strands=CURRICULUM_STRANDS.get(subject,{}),subject_color=SUBJECT_COLORS.get(subject,"#2e86ab"))

@app.route("/bulk_delete",methods=["POST"])
def bulk_delete():
    if not session.get("teacher"): return redirect(url_for("teacher"))
    subject=request.form.get("subject","")
    if not can_access(subject): return redirect(url_for("dashboard"))
    ids=[int(x) for x in request.form.getlist("selected_ids")]
    save_qs(subject,[q for q in load_qs(subject) if q["id"] not in ids])
    session["pdf_msg"]="ok:Deleted {} questions.".format(len(ids))
    return redirect(url_for("dashboard",subject=subject))

@app.route("/bulk_edit_strand",methods=["POST"])
def bulk_edit_strand():
    if not session.get("teacher"): return redirect(url_for("teacher"))
    subject=request.form.get("subject","")
    if not can_access(subject): return redirect(url_for("dashboard"))
    ids=[int(x) for x in request.form.getlist("selected_ids")]
    strand=request.form.get("bulk_strand","").strip(); sub_strand=request.form.get("bulk_sub_strand","").strip()
    classes=request.form.getlist("bulk_classes"); qs=load_qs(subject); updated=0
    for q in qs:
        if q["id"] in ids:
            if strand: q["strand"]=strand; q["sub_strand"]=sub_strand
            if classes: q["assigned_classes"]=classes
            updated+=1
    save_qs(subject,qs); session["pdf_msg"]="ok:Updated {} questions.".format(updated)
    return redirect(url_for("dashboard",subject=subject))

@app.route("/delete_question/<subject>/<int:qid>")
def delete_question(subject,qid):
    if not session.get("teacher"): return redirect(url_for("teacher"))
    if not can_access(subject): return redirect(url_for("dashboard"))
    save_qs(subject,[q for q in load_qs(subject) if q["id"]!=qid])
    return redirect(url_for("dashboard",subject=subject))

@app.route("/upload_pdf",methods=["POST"])
def upload_pdf():
    if not session.get("teacher"): return redirect(url_for("teacher"))
    subject=request.form.get("subject","ICT")
    if not can_access(subject): return redirect(url_for("dashboard"))
    if "pdf" not in request.files: return redirect(url_for("dashboard",subject=subject))
    pdf=request.files["pdf"]
    if not pdf or not pdf.filename or not allowed_pdf(pdf.filename): return redirect(url_for("dashboard",subject=subject))
    fp=os.path.join(UPLOAD_FOLDER,secure_filename(pdf.filename)); pdf.save(fp)
    extracted=extract_pdf(fp)
    if extracted:
        qs=load_qs(subject); mid=max((q["id"] for q in qs),default=0)
        for q in extracted: q["id"]=mid+q["id"]
        qs.extend(extracted); save_qs(subject,qs)
        session["pdf_msg"]="ok:Extracted {} questions.".format(len(extracted))
    else:
        session["pdf_msg"]="err:No questions found. PDF must be typed text not a scanned image."
    return redirect(url_for("dashboard",subject=subject))

@app.route("/delete_result/<int:idx>")
def delete_result(idx):
    if not session.get("teacher"): return redirect(url_for("teacher"))
    rs=load_results()
    if 0<=idx<len(rs):
        t=current_teacher()
        if t.get("is_head") or rs[idx].get("subject","") in t.get("subjects",[]):
            db_delete_result(idx)
    return redirect(url_for("dashboard"))

@app.route("/question_bank")
def question_bank():
    if not session.get("teacher"): return redirect(url_for("teacher"))
    t=current_teacher()
    my_subjects=SUBJECTS if t.get("is_head") else (t.get("subjects") or [])
    sel_sub=request.args.get("subject",my_subjects[0] if my_subjects else "")
    if sel_sub not in my_subjects and my_subjects: sel_sub=my_subjects[0]
    try:
        bank_qs=db_load_bank(sel_sub) if sel_sub else []
    except Exception:
        bank_qs=[]
    return render_template("question_bank.html",bank_qs=bank_qs,selected_subject=sel_sub,
        subjects=my_subjects,school=SCHOOL_NAME,subject_colors=SUBJECT_COLORS,
        subject_icons=SUBJECT_ICONS,classes=CLASSES,strands=CURRICULUM_STRANDS.get(sel_sub,{}))

@app.route("/bank_add",methods=["POST"])
def bank_add():
    if not session.get("teacher"): return redirect(url_for("teacher"))
    subject=request.form.get("subject","")
    if not can_access(subject): return redirect(url_for("question_bank"))
    qtype=request.form.get("type","mcq"); question=request.form.get("question","").strip()
    marks=int(request.form.get("marks",2))
    q={"type":qtype,"question":question,"marks":marks,"image":"","assigned_classes":[],"strand":request.form.get("strand",""),"sub_strand":request.form.get("sub_strand","")}
    if qtype=="mcq": q["options"]=[request.form.get("opt{}".format(i),"").strip() for i in range(1,5)]; q["answer"]=request.form.get("answer","").strip()
    elif qtype=="tf": q["answer"]=request.form.get("tf_answer","True")
    elif qtype=="fitb": q["answer"]=request.form.get("fitb_answer","").strip()
    elif qtype=="theory": q["model_answer"]=request.form.get("model_answer","").strip()
    try: db_add_to_bank(subject,q)
    except Exception as e: print("bank_add error:",e)
    return redirect(url_for("question_bank",subject=subject))

@app.route("/bank_delete/<int:bank_id>")
def bank_delete(bank_id):
    if not session.get("teacher"): return redirect(url_for("teacher"))
    try: db_delete_from_bank(bank_id)
    except Exception as e: print("bank_delete error:",e)
    return redirect(url_for("question_bank"))

@app.route("/bank_to_quiz",methods=["POST"])
def bank_to_quiz():
    if not session.get("teacher"): return redirect(url_for("teacher"))
    subject=request.form.get("subject","")
    if not can_access(subject): return redirect(url_for("question_bank"))
    try:
        ids=[int(x) for x in request.form.getlist("selected_bank_ids")]
        bank_qs=db_load_bank(subject)
        selected=[q for q in bank_qs if q.get("_bank_id") in ids]
        if selected:
            existing=load_qs(subject); mid=max((q["id"] for q in existing),default=0); added=0
            for q in selected:
                mid+=1; nq=dict(q); nq["id"]=mid
                nq.pop("_bank_id",None); nq.pop("_subject",None)
                existing.append(nq); added+=1
            save_qs(subject,existing)
            session["pdf_msg"]="ok:Added {} questions from bank.".format(added)
    except Exception as e:
        print("bank_to_quiz error:",e)
        session["pdf_msg"]="err:Could not add questions from bank."
    return redirect(url_for("dashboard",subject=subject))

@app.route("/ai_generator")
def ai_generator():
    if not session.get("teacher"): return redirect(url_for("teacher"))
    t=current_teacher(); my_subjects=SUBJECTS if t.get("is_head") else (t.get("subjects") or [])
    sel_sub=request.args.get("subject",my_subjects[0] if my_subjects else "")
    return render_template("ai_generator.html",school=SCHOOL_NAME,subjects=my_subjects,
        selected_subject=sel_sub,subject_icons=SUBJECT_ICONS,classes=CLASSES,
        strands=CURRICULUM_STRANDS.get(sel_sub,{}),curriculum_strands=CURRICULUM_STRANDS)

@app.route("/ai_generate_questions",methods=["POST"])
def ai_generate_questions():
    if not session.get("teacher"): return jsonify({"error":"Not logged in"}),401
    if not REQUESTS_OK: return jsonify({"error":"Run: pip install requests"})
    if not GROQ_API_KEY: return jsonify({"error":"GROQ_API_KEY not set in Railway environment variables."})
    data=request.json; topic=data.get("topic","").strip(); subject=data.get("subject","ICT")
    strand=data.get("strand",""); sub_strand=data.get("sub_strand","")
    cls=data.get("class_level","JHS 1"); count=min(int(data.get("count",10)),20)
    difficulty=data.get("difficulty","intermediate"); types=data.get("types",["mcq","tf","fitb"])
    if not topic: return jsonify({"error":"Please enter a topic."})
    diff_map={"basic":"simple for Primary school","intermediate":"appropriate for JHS","challenging":"BECE examination standard"}
    type_map={"mcq":"Multiple Choice (4 options)","tf":"True/False","fitb":"Fill in the Blank (use _______)","theory":"Theory/Essay with model answer"}
    prompt="""You are a Ghana Basic School teacher. Generate exactly {count} exam questions.
Subject: {subject} | Class: {cls} | Topic: {topic}
Strand: {strand} | Difficulty: {diff}
Question types: {types}
IMPORTANT FOR MCQ: The "answer" field MUST be exactly "A","B","C", or "D".
Return ONLY valid JSON:
{{"questions":[
  {{"type":"mcq","question":"text","options":["option1","option2","option3","option4"],"answer":"A","marks":2}},
  {{"type":"tf","question":"text","answer":"True","marks":1}},
  {{"type":"fitb","question":"The _______ is...","answer":"word","marks":2}},
  {{"type":"theory","question":"Explain... [3 marks]","model_answer":"key points","marks":3}}
]}}""".format(count=count,subject=subject,cls=cls,topic=topic,strand="{}-{}".format(strand,sub_strand) if strand else "General",diff=diff_map.get(difficulty,"appropriate for JHS"),types=", ".join([type_map.get(t,t) for t in types]))
    try:
        resp=req_lib.post(GROQ_URL,headers={"Authorization":"Bearer "+GROQ_API_KEY,"Content-Type":"application/json"},json={"model":GROQ_MODEL,"max_tokens":4000,"temperature":0.7,"messages":[{"role":"system","content":"You are a Ghana teacher. Respond with valid JSON only."},{"role":"user","content":prompt}]},timeout=60)
        if resp.status_code!=200: return jsonify({"error":"Groq error {}: {}".format(resp.status_code,resp.text[:120])})
        ai_text=resp.json()["choices"][0]["message"]["content"].strip()
        if "```json" in ai_text: ai_text=ai_text.split("```json")[1].split("```")[0].strip()
        elif "```" in ai_text: ai_text=ai_text.split("```")[1].split("```")[0].strip()
        fb=ai_text.find("{"); lb=ai_text.rfind("}")
        if fb!=-1 and lb!=-1: ai_text=ai_text[fb:lb+1]
        questions=json.loads(ai_text).get("questions",[])
        clean_qs=[]
        for q in questions:
            qtype=q.get("type","mcq")
            if qtype not in ["mcq","tf","fitb","theory"]: continue
            clean={"type":qtype,"question":q.get("question","").strip(),"marks":int(q.get("marks",1 if qtype=="tf" else 2)),"image":"","assigned_classes":[],"strand":strand,"sub_strand":sub_strand}
            if qtype=="mcq":
                opts=q.get("options",[])
                if len(opts)<2: continue
                clean["options"]=opts[:4]; raw_ans=q.get("answer","").strip()
                if raw_ans.upper() in ["A","B","C","D"]: clean["answer"]=raw_ans.upper()
                else:
                    matched_letter="A"
                    for i,opt in enumerate(opts[:4]):
                        if opt.strip().lower()==raw_ans.lower(): matched_letter=["A","B","C","D"][i]; break
                    clean["answer"]=matched_letter
            elif qtype=="tf":
                ans=str(q.get("answer","True")).strip().capitalize()
                clean["answer"]=ans if ans in ["True","False"] else "True"
            elif qtype=="fitb": clean["answer"]=q.get("answer","").strip()
            elif qtype=="theory": clean["model_answer"]=q.get("model_answer","").strip(); clean["answer"]=clean["model_answer"]
            if clean.get("question"): clean_qs.append(clean)
        return jsonify({"questions":clean_qs})
    except json.JSONDecodeError: return jsonify({"error":"AI returned invalid format. Try again."})
    except Exception as e:
        import traceback; print(traceback.format_exc())
        return jsonify({"error":"AI failed: "+str(e)[:150]})

@app.route("/ai_add_questions",methods=["POST"])
def ai_add_questions():
    if not session.get("teacher"): return jsonify({"error":"Not logged in"}),401
    data=request.json; subject=data.get("subject","ICT")
    if not can_access(subject): return jsonify({"error":"No access"}),403
    strand=data.get("strand",""); sub_strand=data.get("sub_strand","")
    assigned_classes=data.get("assigned_classes",[]); questions=data.get("questions",[]); existing=load_qs(subject)
    mid=max((q["id"] for q in existing),default=0); added=0
    for q in questions:
        mid+=1; q["id"]=mid; q["strand"]=strand or q.get("strand",""); q["sub_strand"]=sub_strand or q.get("sub_strand","")
        q["image"]=""; q["assigned_classes"]=assigned_classes
        if "answer" not in q and q.get("type")=="theory": q["answer"]=q.get("model_answer","")
        existing.append(q); added+=1
    save_qs(subject,existing); return jsonify({"ok":True,"added":added})

@app.route("/student_report")
def student_report():
    if not session.get("teacher"): return redirect(url_for("teacher"))
    t=current_teacher(); results=load_results()
    if not t.get("is_head"): results=[r for r in results if r.get("subject","") in (t.get("subjects") or [])]
    students={}
    for r in results:
        key=r["name"].strip().lower()+"|"+r["class"].strip().lower()
        if key not in students: students[key]={"name":r["name"],"class":r["class"],"tests":[]}
        students[key]["tests"].append(r)
    summaries=[]
    for st in students.values():
        avg=round(sum(x["percentage"] for x in st["tests"])/len(st["tests"]),1)
        g,rem=get_grade(avg)
        summaries.append({"name":st["name"],"class":st["class"],"tests":st["tests"],"avg":avg,"grade":g,"remark":rem,"count":len(st["tests"])})
    summaries.sort(key=lambda x:(-x["avg"],x["name"]))
    return render_template("student_report.html",summaries=summaries,school=SCHOOL_NAME,subject_icons=SUBJECT_ICONS,grade_color=gcolor)

@app.route("/student_report_word/<path:key>")
def student_report_word(key):
    if not session.get("teacher"): return redirect(url_for("teacher"))
    name,cls=key.split("|",1)
    tests=[r for r in load_results() if r["name"].strip().lower()==name and r["class"].strip().lower()==cls]
    if not tests: return redirect(url_for("student_report"))
    avg=round(sum(t["percentage"] for t in tests)/len(tests),1); grade,remark=get_grade(avg)
    doc=Document()
    for sec in doc.sections: sec.top_margin=sec.bottom_margin=Inches(0.8); sec.left_margin=sec.right_margin=Inches(1.0)
    h=doc.add_heading(SCHOOL_NAME,0); h.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for r in h.runs: r.font.color.rgb=RGBColor(0x1A,0x3A,0x5C); r.font.size=Pt(18)
    s=doc.add_paragraph("Student Academic Performance Report"); s.alignment=WD_ALIGN_PARAGRAPH.CENTER
    s.runs[0].font.size=Pt(13); s.runs[0].font.color.rgb=RGBColor(0x2E,0x86,0xAB); doc.add_paragraph()
    tbl=doc.add_table(rows=3,cols=2); tbl.style="Table Grid"
    for i,(l,v) in enumerate([("Student Name:",tests[0]["name"]),("Class:",tests[0]["class"]),("Date:",datetime.datetime.now().strftime("%d/%m/%Y"))]):
        tbl.rows[i].cells[0].text=l; tbl.rows[i].cells[1].text=v; tbl.rows[i].cells[0].paragraphs[0].runs[0].font.bold=True
    doc.add_paragraph()
    grgb={"A":RGBColor(0x27,0xAE,0x60),"B":RGBColor(0x29,0x80,0xB9),"C":RGBColor(0x8E,0x44,0xAD),"D":RGBColor(0xF3,0x9C,0x12),"E":RGBColor(0xE6,0x7E,0x22),"F":RGBColor(0xE7,0x4C,0x3C)}
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run=p.add_run("OVERALL: {}% | GRADE: {} | {}".format(avg,grade,remark.upper()))
    run.font.bold=True; run.font.size=Pt(14); run.font.color.rgb=grgb.get(grade,RGBColor(0,0,0))
    doc.add_paragraph(); doc.add_heading("Subject Results",level=2)
    t2=doc.add_table(rows=1,cols=5); t2.style="Table Grid"
    for i,h in enumerate(["#","Subject","Assessment","Score","Grade"]):
        t2.rows[0].cells[i].text=h; t2.rows[0].cells[i].paragraphs[0].runs[0].font.bold=True
    for i,r in enumerate(tests,1):
        row=t2.add_row(); row.cells[0].text=str(i); row.cells[1].text=r.get("subject","")
        row.cells[2].text=r.get("assessment_label",""); row.cells[3].text="{}/{}".format(r["score"],r["total"]); row.cells[4].text=r["grade"]
    doc.add_paragraph(); doc.add_heading("Teacher Remark",level=2)
    rem=doc.add_table(rows=1,cols=1); rem.style="Table Grid"; rem.rows[0].cells[0].text="\n\n\n"
    doc.add_paragraph(); sig=doc.add_paragraph(); sig.alignment=WD_ALIGN_PARAGRAPH.CENTER
    sig.add_run("Class Teacher: ________________ Date: ________________").font.size=Pt(11)
    fname="report_{}_all.docx".format(tests[0]["name"].replace(" ","_")); doc.save(fname)
    return send_file(fname,as_attachment=True)

@app.route("/export_excel")
def export_excel():
    if not session.get("teacher"): return redirect(url_for("teacher"))
    t=current_teacher(); results=load_results()
    if not t.get("is_head"): results=[r for r in results if r.get("subject","") in (t.get("subjects") or [])]
    wb=Workbook(); ws=wb.active; ws.title="Results"
    ws.merge_cells("A1:K1"); ws["A1"]=SCHOOL_NAME+" — Quiz Results"
    ws["A1"].font=Font(bold=True,size=14,color="FFFFFF"); ws["A1"].fill=PatternFill("solid",fgColor="1A3A5C"); ws["A1"].alignment=Alignment(horizontal="center")
    ws.merge_cells("A2:K2"); ws["A2"]="Generated: "+datetime.datetime.now().strftime("%d/%m/%Y %H:%M")
    ws["A2"].font=Font(italic=True,size=11); ws["A2"].alignment=Alignment(horizontal="center")
    hdrs=["#","Name","Class","Subject","Assessment","Strand","Score","Total","%","Grade","Date"]
    for col,h in enumerate(hdrs,1):
        cell=ws.cell(row=4,column=col,value=h); cell.font=Font(bold=True,color="FFFFFF",size=11)
        cell.fill=PatternFill("solid",fgColor="2E86AB"); cell.alignment=Alignment(horizontal="center")
    gcols={"A":"D5F5E3","B":"D6EAF8","C":"E8DAEF","D":"FDEBD0","E":"FAE5D3","F":"FADBD8"}
    thin=Side(style="thin",color="CCCCCC")
    for i,r in enumerate(results,1):
        row=4+i
        data=[i,r["name"],r["class"],r.get("subject",""),r.get("assessment_label",""),r.get("strand",""),r["score"],r["total"],"{}%".format(r["percentage"]),r["grade"],r.get("date","")]
        for col,val in enumerate(data,1):
            cell=ws.cell(row=row,column=col,value=val); cell.alignment=Alignment(horizontal="center")
            cell.fill=PatternFill("solid",fgColor=gcols.get(r["grade"],"FFFFFF")); cell.border=Border(left=thin,right=thin,top=thin,bottom=thin)
    for letter,w in zip("ABCDEFGHIJK",[5,20,10,20,12,16,7,7,8,7,12]): ws.column_dimensions[letter].width=w
    fname="results_{}.xlsx".format(datetime.datetime.now().strftime("%Y%m%d_%H%M")); wb.save(fname); return send_file(fname,as_attachment=True)

@app.route("/export_word/<int:idx>")
def export_word(idx):
    if not session.get("teacher"): return redirect(url_for("teacher"))
    rs=load_results()
    if idx>=len(rs): return redirect(url_for("dashboard"))
    r=rs[idx]; doc=Document()
    for sec in doc.sections: sec.top_margin=sec.bottom_margin=Inches(0.8); sec.left_margin=sec.right_margin=Inches(1.0)
    h=doc.add_heading(SCHOOL_NAME,0); h.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs: run.font.color.rgb=RGBColor(0x1A,0x3A,0x5C); run.font.size=Pt(18)
    s=doc.add_paragraph("{} — {} Report".format(r.get("subject",""),r.get("assessment_label",""))); s.alignment=WD_ALIGN_PARAGRAPH.CENTER
    s.runs[0].font.size=Pt(13); s.runs[0].font.color.rgb=RGBColor(0x2E,0x86,0xAB); doc.add_paragraph()
    tbl=doc.add_table(rows=5,cols=2); tbl.style="Table Grid"
    for i,(l,v) in enumerate([("Name:",r["name"]),("Class:",r["class"]),("Subject:",r.get("subject","")),("Assessment:",r.get("assessment_label","")),("Date:",r.get("date",""))]):
        tbl.rows[i].cells[0].text=l; tbl.rows[i].cells[1].text=v; tbl.rows[i].cells[0].paragraphs[0].runs[0].font.bold=True
    doc.add_paragraph()
    grgb={"A":RGBColor(0x27,0xAE,0x60),"B":RGBColor(0x29,0x80,0xB9),"C":RGBColor(0x8E,0x44,0xAD),"D":RGBColor(0xF3,0x9C,0x12),"E":RGBColor(0xE6,0x7E,0x22),"F":RGBColor(0xE7,0x4C,0x3C)}
    sp=doc.add_paragraph(); sp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run=sp.add_run("SCORE: {}/{} | {}% | GRADE: {} | {}".format(r["score"],r["total"],r["percentage"],r["grade"],r["remark"].upper()))
    run.font.bold=True; run.font.size=Pt(13); run.font.color.rgb=grgb.get(r["grade"],RGBColor(0,0,0))
    doc.add_paragraph(); doc.add_heading("Detailed Results",level=2)
    dtbl=doc.add_table(rows=1,cols=5); dtbl.style="Table Grid"
    for i,h in enumerate(["#","Question","Your Answer","Correct","Marks"]):
        dtbl.rows[0].cells[i].text=h; dtbl.rows[0].cells[i].paragraphs[0].runs[0].font.bold=True
    for i,d in enumerate(r["details"],1):
        row=dtbl.add_row(); row.cells[0].text=str(i); row.cells[1].text=d["question"][:80]
        row.cells[2].text=d["given"][:60]; row.cells[3].text=d["correct"][:60] if d["type"]!="theory" else "(model answer)"
        row.cells[4].text="{}/{}".format(d["marks_earned"],d["marks_total"])
        col=RGBColor(0x27,0xAE,0x60) if d["correct_flag"] else RGBColor(0xE7,0x4C,0x3C)
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs: run.font.color.rgb=col
    doc.add_paragraph()
    ft=doc.add_paragraph("Generated by {} — {}".format(SCHOOL_NAME,datetime.datetime.now().strftime("%d/%m/%Y %H:%M")))
    ft.alignment=WD_ALIGN_PARAGRAPH.CENTER; ft.runs[0].font.size=Pt(9); ft.runs[0].font.italic=True
    fname="report_{}.docx".format(r["name"].replace(" ","_")); doc.save(fname); return send_file(fname,as_attachment=True)

@app.route("/pdf_template")
def pdf_template():
    if not session.get("teacher"): return redirect(url_for("teacher"))
    doc=Document()
    for sec in doc.sections: sec.top_margin=sec.bottom_margin=Inches(1.0); sec.left_margin=sec.right_margin=Inches(1.2)
    h=doc.add_heading("PDF Question Template — "+SCHOOL_NAME,0); h.alignment=WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("IMPORTANT: Type in Microsoft Word and Save As PDF. Scanned images will NOT work."); doc.add_paragraph()
    for title,body in [("MCQ","1. What does CPU stand for?\nA. Central Processing Unit\nB. Computer Processing Unit\nC. Central Program Unit\nD. Computer Program Unit\nAnswer: Central Processing Unit"),("True/False","TF: 1. A monitor is an output device.\nAnswer: True"),("Fill in Blank","FB: 1. The brain of the computer is called the _______.\nAnswer: CPU"),("Theory","TH: 1. Explain what a computer network is. [5 marks]\nModel: A computer network connects computers to share resources and data.")]:
        doc.add_heading(title,2); doc.add_paragraph(body); doc.add_paragraph()
    fname="pdf_template.docx"; doc.save(fname); return send_file(fname,as_attachment=True)

@app.route("/backup_results")
def backup_results():
    if not session.get("teacher"): return redirect(url_for("teacher"))
    ts=datetime.datetime.now().strftime("%Y%m%d_%H%M"); zf="backup_{}.zip".format(ts)
    with zipfile.ZipFile(zf,"w") as z:
        if os.path.exists(RESULTS_FILE): z.write(RESULTS_FILE)
        for f in os.listdir("."):
            if (f.startswith("questions_") or f.startswith("bece_")) and f.endswith(".json"): z.write(f)
    return send_file(zf,as_attachment=True)

@app.route("/restore_results",methods=["POST"])
def restore_results():
    if not session.get("teacher"): return redirect(url_for("teacher"))
    f=request.files.get("restore_file")
    if not f or not f.filename: return redirect(url_for("dashboard"))
    fname=f.filename.lower()
    if fname.endswith(".json"):
        with _lock: json.dump(json.loads(f.read().decode("utf-8")),open(RESULTS_FILE,"w"),indent=2)
        session["restore_msg"]="ok:Results restored."
    elif fname.endswith(".zip"):
        import io; zf=zipfile.ZipFile(io.BytesIO(f.read()))
        for name in zf.namelist():
            if name==RESULTS_FILE or name.startswith("questions_") or name.startswith("bece_"): open(name,"wb").write(zf.read(name))
        session["restore_msg"]="ok:Data restored from ZIP."
    return redirect(url_for("dashboard"))

@app.route("/list_backups")
def list_backups():
    if not session.get("teacher"): return redirect(url_for("teacher"))
    backups=[]
    if os.path.exists("backups"):
        for f in sorted(os.listdir("backups"),reverse=True):
            if f.endswith(".json"):
                fp=os.path.join("backups",f); count=0
                try: count=len(json.load(open(fp)))
                except: pass
                backups.append({"filename":f,"date":datetime.datetime.fromtimestamp(os.path.getmtime(fp)).strftime("%d/%m/%Y %H:%M"),"size":round(os.path.getsize(fp)/1024,1),"count":count})
    return render_template("backups.html",backups=backups,school=SCHOOL_NAME)

@app.route("/restore_backup/<filename>")
def restore_backup(filename):
    if not session.get("teacher"): return redirect(url_for("teacher"))
    fp=os.path.join("backups",filename)
    if os.path.exists(fp):
        with _lock: shutil.copy(fp,RESULTS_FILE)
    return redirect(url_for("dashboard"))

@app.route("/download_backup/<filename>")
def download_backup(filename):
    if not session.get("teacher"): return redirect(url_for("teacher"))
    fp=os.path.join("backups",filename)
    if os.path.exists(fp): return send_file(fp,as_attachment=True)
    return redirect(url_for("list_backups"))

@app.route("/bece")
def bece_home():
    sy={}
    for sub in BECE_SUBJECTS:
        yrs=bece_years_for(sub)
        if yrs: sy[sub]=yrs
    return render_template("bece_home.html",school=SCHOOL_NAME,subject_years=sy,
        bece_subjects=BECE_SUBJECTS,subject_icons=SUBJECT_ICONS,subject_colors=SUBJECT_COLORS)

@app.route("/bece/start",methods=["POST"])
def bece_start():
    name=request.form.get("name","").strip(); subject=request.form.get("subject","").strip()
    year=request.form.get("year","").strip(); mode=request.form.get("mode","practice")
    if not name or not subject or not year: return redirect(url_for("bece_home"))
    qs=load_bece(subject,year)
    if not qs: return redirect(url_for("bece_home"))
    if mode=="exam": qs=qs.copy(); random.shuffle(qs)
    session.update({"bece_name":name,"bece_subject":subject,"bece_year":year,"bece_mode":mode,"bece_questions":qs})
    return render_template("bece_quiz.html",questions=qs,student_name=name,subject=subject,year=year,mode=mode,
        subject_color=SUBJECT_COLORS.get(subject,"#c0392b"),subject_icon=SUBJECT_ICONS.get(subject,"📝"),
        timer=30*60 if mode=="exam" else 0,school=SCHOOL_NAME)

@app.route("/bece/submit",methods=["POST"])
def bece_submit():
    qs=session.get("bece_questions",[]); name=session.get("bece_name","Student")
    subject=session.get("bece_subject",""); year=session.get("bece_year",""); mode=session.get("bece_mode","practice")
    total=sum(q.get("marks",1) for q in qs); earned=0; details=[]
    for q in qs:
        qid=str(q["id"]); given=request.form.get("q{}".format(qid),"").strip()
        correct=q.get("answer","").strip(); marks=q.get("marks",1)
        if q["type"] in ["mcq","tf","fitb"]:
            if q["type"]=="mcq":
                opts=q.get("options",[]); letters=["A","B","C","D"]
                given_letter=given.upper() if given.upper() in letters else ""
                if not given_letter:
                    for i,opt in enumerate(opts):
                        if opt.strip().lower()==given.strip().lower() and i<4: given_letter=letters[i]; break
                correct_letter=correct.upper() if correct.upper() in letters else ""
                if not correct_letter:
                    for i,opt in enumerate(opts):
                        if opt.strip().lower()==correct.strip().lower() and i<4: correct_letter=letters[i]; break
                right=given_letter==correct_letter if given_letter and correct_letter else False
                given_disp=("{}: {}".format(given_letter,opts[letters.index(given_letter)]) if given_letter in letters and letters.index(given_letter)<len(opts) else given_letter) if given_letter else "Not answered"
                correct_disp=("{}: {}".format(correct_letter,opts[letters.index(correct_letter)]) if correct_letter in letters and letters.index(correct_letter)<len(opts) else correct_letter) if correct_letter else correct
            else:
                right=given.lower()==correct.lower(); given_disp=given or "Not answered"; correct_disp=correct
            got=marks if right else 0; earned+=got
            details.append({"question":q["question"],"type":q["type"],"given":given_disp,"correct":correct_disp,"correct_flag":right,"marks_earned":got,"marks_total":marks})
        elif q["type"]=="theory":
            got,matched=keyword_score(given,q.get("model_answer",""),marks); earned+=got
            details.append({"question":q["question"],"type":"theory","given":given or "Not answered","correct":q.get("model_answer",""),"correct_flag":got>=marks*0.5,"marks_earned":got,"marks_total":marks,"matched_keywords":matched})
    pct=round((earned/total)*100,1) if total>0 else 0
    grade,remark=get_grade(pct)
    result={"name":name,"subject":subject,"year":year,"mode":mode,"score":earned,"total":total,"percentage":pct,"grade":grade,"remark":remark,"date":datetime.datetime.now().strftime("%d/%m/%Y"),"details":details}
    return render_template("bece_result.html",result=result,school=SCHOOL_NAME,
        grade_color=gcolor(grade),subject_color=SUBJECT_COLORS.get(subject,"#c0392b"),
        subject_icon=SUBJECT_ICONS.get(subject,"📝"),show_answers=mode=="practice")

@app.route("/superadmin", methods=["GET","POST"])
def superadmin():
    error = None
    if request.method == "POST":
        pw = request.form.get("password", "").strip()
        admin_pw = os.environ.get("ADMIN_PASSWORD", "admin2024")
        secret_pw = app.secret_key
        if pw == admin_pw or pw == secret_pw:
            session["superadmin"] = True
            return redirect(url_for("superadmin"))
        error = "Wrong password. Please try again."
    if not session.get("superadmin"):
        return render_template("superadmin.html", school=SCHOOL_NAME,
            logged_in=False, error=error,
            pending={}, approved={}, rejected={}, schools=[], subjects=SUBJECTS, classes=CLASSES)
    # Logged in — load all data
    fresh = load_custom_teachers()
    CUSTOM_TEACHERS.update(fresh)
    pending  = {k:v for k,v in CUSTOM_TEACHERS.items() if v.get("status") == "pending"}
    approved = {k:v for k,v in CUSTOM_TEACHERS.items() if v.get("status") == "approved"}
    rejected = {k:v for k,v in CUSTOM_TEACHERS.items() if v.get("status") == "rejected"}
    try:
        schools = db_get_all_schools()
    except Exception:
        schools = []
    return render_template("superadmin.html", school=SCHOOL_NAME,
        logged_in=True, error=None,
        pending=pending, approved=approved, rejected=rejected,
        schools=schools, subjects=SUBJECTS, classes=CLASSES)

@app.route("/superadmin/approve/<uid>")
def approve_teacher(uid):
    if not session.get("superadmin"): return redirect(url_for("superadmin"))
    data=load_custom_teachers()
    if uid in data: data[uid]["status"]="approved"; save_custom_teachers(data); CUSTOM_TEACHERS.update(data)
    return redirect(url_for("superadmin"))

@app.route("/superadmin/reject/<uid>")
def reject_teacher(uid):
    if not session.get("superadmin"): return redirect(url_for("superadmin"))
    data=load_custom_teachers()
    if uid in data: data[uid]["status"]="rejected"; save_custom_teachers(data); CUSTOM_TEACHERS.update(data)
    return redirect(url_for("superadmin"))

@app.route("/superadmin/delete/<uid>")
def delete_teacher(uid):
    if not session.get("superadmin"): return redirect(url_for("superadmin"))
    data=load_custom_teachers()
    if uid in data: del data[uid]; save_custom_teachers(data); CUSTOM_TEACHERS.clear(); CUSTOM_TEACHERS.update(data)
    return redirect(url_for("superadmin"))

@app.route("/superadmin/edit_teacher/<uid>",methods=["POST"])
def edit_teacher(uid):
    if not session.get("superadmin"): return redirect(url_for("superadmin"))
    data=load_custom_teachers()
    if uid in data:
        new_subjects=request.form.getlist("subjects"); new_class=request.form.get("managed_class","").strip(); new_role=request.form.get("role","subject_teacher").strip()
        if new_subjects: data[uid]["subjects"]=new_subjects; data[uid]["managed_class"]=new_class if new_class else None; data[uid]["role"]=new_role; save_custom_teachers(data); CUSTOM_TEACHERS.update(data)
    return redirect(url_for("superadmin"))

@app.route("/superadmin/logout")
def superadmin_logout():
    session.pop("superadmin",None); return redirect(url_for("superadmin"))

@app.route("/register_school",methods=["GET","POST"])
def register_school():
    if request.method=="POST":
        data={"school_name":request.form.get("school_name","").strip(),"school_code":request.form.get("school_code","").strip().upper(),"head_name":request.form.get("head_name","").strip(),"head_email":request.form.get("head_email","").strip(),"head_phone":request.form.get("head_phone","").strip(),"head_password":generate_password_hash(request.form.get("head_password","")),"region":request.form.get("region","").strip()}
        if not all([data["school_name"],data["school_code"],data["head_name"],data["head_password"]]):
            return render_template("register_school.html",school=SCHOOL_NAME,regions=GHANA_REGIONS,error="Please fill in all required fields.")
        try:
            if db_school_code_exists(data["school_code"]): return render_template("register_school.html",school=SCHOOL_NAME,regions=GHANA_REGIONS,error="That school code is already registered.")
            db_register_school(data)
        except Exception as e:
            print("register_school error:",e)
            return render_template("register_school.html",school=SCHOOL_NAME,regions=GHANA_REGIONS,error="Registration failed. Please try again.")
        return redirect(url_for("landing"))
    return render_template("register_school.html",school=SCHOOL_NAME,regions=GHANA_REGIONS,error=None)

if __name__=="__main__":
    app.run(debug=True,host="0.0.0.0",port=5000)
